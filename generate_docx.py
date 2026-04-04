from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Project Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Proposed System:', level=1)
    doc.add_paragraph('Web-Based AI-Assisted Student Information and Academic Monitoring System for Calvary Christian Academy', style='Subtitle')
    
    # 1. Techstack
    doc.add_heading('1. Techstack', level=1)
    doc.add_paragraph('The system utilizes a modern full-stack architecture with a focus on AI integration and responsive design.', style='Body Text')
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    
    tech_data = [
        ('Frontend Framework', 'React 19 (Vite)'),
        ('Styling', 'TailwindCSS 4'),
        ('Backend Framework', 'FastAPI (Python)'),
        ('Database', 'SQLite (Development) / PostgreSQL (Optional)'),
        ('ORM', 'SQLAlchemy'),
        ('AI/OCR', 'PyTesseract (Tesseract OCR), Pillow'),
        ('Machine Learning', 'Scikit-learn, Pandas, NumPy'),
        ('API Documentation', 'Swagger UI / Redoc (Built-in FastAPI)')
    ]
    
    for comp, tech in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = tech

    # 2. Software Applications Used
    doc.add_heading('2. Software Applications Used', level=1)
    sw_apps = [
        ('Visual Studio Code', 'Current Stable (IDE)'),
        ('Python', '3.10+ (Backend Runtime)'),
        ('Node.js', '18+ (Frontend Runtime)'),
        ('Git', 'Version Control System'),
        ('Postman', 'API Testing and Documentation'),
        ('Tesseract OCR', 'v5.x (OCR Engine)'),
        ('Vite', 'Frontend Build Tool'),
        ('Microsoft Edge / Chrome', 'Modern Web Browsers')
    ]
    
    sw_table = doc.add_table(rows=1, cols=2)
    sw_table.style = 'Table Grid'
    hdr_cells = sw_table.rows[0].cells
    hdr_cells[0].text = 'Application'
    hdr_cells[1].text = 'Version / Role'
    
    for app, ver in sw_apps:
        row_cells = sw_table.add_row().cells
        row_cells[0].text = app
        row_cells[1].text = ver

    # 3. Hardware Requirements
    doc.add_heading('3. Hardware Requirements', level=1)
    doc.add_paragraph('Minimum requirements for administrative and staff usage.')
    
    hw_data = [
        ('Computing Device Type', 'PC or Laptop (for clinic staff use)'),
        ('Processor', 'Intel Core i3 or Ryzen 3 (Dual-core 2.4GHz+)'),
        ('Memory (RAM)', '8GB DDR4 (Minimum)'),
        ('Storage', '256GB SSD (for fast system performance)'),
        ('Network Connectivity', 'Broadband Modem/Router with stable internet connection'),
        ('Scanner/Peripherals', 'High-resolution scanner or camera (for OCR uploads)'),
        ('Display', 'Standard 1080p monitor for dashboard visibility')
    ]
    
    hw_table = doc.add_table(rows=1, cols=2)
    hw_table.style = 'Table Grid'
    hdr_cells = hw_table.rows[0].cells
    hdr_cells[0].text = 'Requirement'
    hdr_cells[1].text = 'Specification'
    
    for req, spec in hw_data:
        row_cells = hw_table.add_row().cells
        row_cells[0].text = req
        row_cells[1].text = spec

    # 4. SDLC Model and Timetable
    doc.add_heading('4. SDLC Model: Agile Model', level=1)
    doc.add_paragraph('The project follows the Agile SDLC model to allow for iterative development, continuous feedback, and flexibility in AI feature refinement.')
    
    doc.add_heading('Proposed Timetable (6 Months)', level=2)
    timetable = [
        ('Month 1', 'Planning & Requirements Gathering', 'Identifying system problems, objectives, and user needs.'),
        ('Month 2', 'System Design', 'Architecture design, UI/UX prototyping, and AI logic planning.'),
        ('Month 3', 'Core Development', 'Building Student Info Module and Database structure.'),
        ('Month 4', 'AI Integration', 'OCR automation development and ML tracking implementation.'),
        ('Month 5', 'Testing & Validation', 'System testing, OCR accuracy validation, and bug fixing.'),
        ('Month 6', 'Deployment & Improvement', 'Final deployment, staff training, and feature refinement.')
    ]
    
    tt_table = doc.add_table(rows=1, cols=3)
    tt_table.style = 'Table Grid'
    hdr_cells = tt_table.rows[0].cells
    hdr_cells[0].text = 'Timeframe'
    hdr_cells[1].text = 'Phase'
    hdr_cells[2].text = 'Key Activities'
    
    for mo, ph, act in timetable:
        row_cells = tt_table.add_row().cells
        row_cells[0].text = mo
        row_cells[1].text = ph
        row_cells[2].text = act

    # 5. Output: Features of Proposed System
    doc.add_heading('5. Features of the Proposed System', level=1)
    features = [
        ('Automated Enrollment (OCR)', 'Extracts student data from uploaded forms directly into the database, reducing repetitive encoding.'),
        ('Centralized Student Database', 'Secure and organized storage for all student, academic, and administrative records.'),
        ('AI-Powered Academic Monitoring', 'Early Warning System (EWS) to detect declining performance trends based on grades.'),
        ('Tuition Payment Risk Prediction', 'Utilizes Machine Learning to classify payment risks and improve financial oversight.'),
        ('Intelligent Report Generator', 'Generates institutional and student reports in seconds instead of hours.'),
        ('Responsive Student Portal', 'Allows students and parents to monitor academic progress and records in real-time.')
    ]
    
    for feat, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(feat + ': ')
        run.bold = True
        p.add_run(desc)

    doc.save('Project_Documentation.docx')
    print("Project_Documentation.docx has been created successfully.")

if __name__ == '__main__':
    create_doc()
