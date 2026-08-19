import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F4E78") # Deep Blue
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)

    # Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F2F4F7" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            for p in row_cells[c_idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(50, 50, 50)

    # Set Column Widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph() # Spacing after table
    return table

def generate_doc():
    doc = Document()

    # Set Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Project Documentation: Calvary Christian Academy")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(31, 78, 120)

    # Proposed System
    p_sub = doc.add_paragraph()
    r_sub_lbl = p_sub.add_run("Proposed System: ")
    r_sub_lbl.bold = True
    r_sub_val = p_sub.add_run("Web-Based AI-Assisted Student Information and Academic Monitoring System")
    p_sub.runs[0].font.size = Pt(11)
    p_sub.runs[1].font.size = Pt(11)

    doc.add_paragraph()

    # 1. Techstack
    h1 = doc.add_heading("1. Technology Stack", level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 120)
    p_stack = doc.add_paragraph("The system utilizes a modern decoupled architecture with a focus on AI integration and responsive design.")
    p_stack.runs[0].font.name = "Calibri"

    # Frontend Table
    doc.add_heading("Frontend Framework & Styling", level=2).runs[0].font.color.rgb = RGBColor(46, 117, 182)
    fe_headers = ["Component", "Technology", "Version / Detail"]
    fe_data = [
        ["Frontend Framework", "React", "^19.2.4 (Vite ^8.0.1)"],
        ["Styling", "Tailwind CSS", "^4.2.2"],
        ["Language", "JavaScript (ES6+)", "With JSX"]
    ]
    create_styled_table(doc, fe_headers, fe_data, col_widths=[2.0, 2.0, 2.5])

    # Backend Table
    doc.add_heading("Backend Framework & Database", level=2).runs[0].font.color.rgb = RGBColor(46, 117, 182)
    be_headers = ["Component", "Technology", "Version / Detail"]
    be_data = [
        ["Backend Framework", "FastAPI", "High-performance Python API framework"],
        ["ASGI Server", "Uvicorn", "[standard] for production-grade serving"],
        ["Database", "SQLite", "Development / Local deployment (.db files)"],
        ["ORM", "SQLAlchemy", "Relational database object mapping"],
        ["Data Validation", "Pydantic", "Schema and type validation"],
        ["Authentication", "Passlib, python-jose", "Password hashing and JWT authentication"],
        ["API Documentation", "Swagger UI / Redoc", "Built-in FastAPI interactive documentation"]
    ]
    create_styled_table(doc, be_headers, be_data, col_widths=[2.0, 2.0, 2.5])

    # AI & ML Table
    doc.add_heading("AI & Machine Learning Integrations", level=2).runs[0].font.color.rgb = RGBColor(46, 117, 182)
    ai_headers = ["Component", "Technology", "Role"]
    ai_data = [
        ["Machine Learning", "Scikit-learn, NumPy, Pandas", "RandomForest & GradientBoosting models for predictive academic and tuition risk tracking."],
        ["Generative AI", "Google Generative AI", "Integration with Gemini (2.0 Flash) for automated narrative report generation."]
    ]
    create_styled_table(doc, ai_headers, ai_data, col_widths=[2.0, 2.0, 2.5])

    # 2. Software Applications Used
    h1_sw = doc.add_heading("2. Software Applications Used", level=1)
    h1_sw.runs[0].font.color.rgb = RGBColor(31, 78, 120)
    sw_headers = ["Application", "Version / Role", "Description"]
    sw_data = [
        ["Visual Studio Code", "Current Stable", "Primary Integrated Development Environment (IDE)."],
        ["Python", "3.10+", "Backend runtime environment and ML execution."],
        ["Node.js", "18.0+", "Frontend runtime environment for Vite and npm."],
        ["Git", "Current Stable", "Version control system."],
        ["Postman", "Latest", "API endpoint testing and payload validation."],
        ["Vite", "^8.0.1", "Frontend build tool and development server."],
        ["Modern Web Browser", "Chrome / Edge / Safari", "Client interface testing and application interaction."]
    ]
    create_styled_table(doc, sw_headers, sw_data, col_widths=[1.8, 1.7, 3.0])

    # 3. Hardware Requirements
    h1_hw = doc.add_heading("3. Hardware Requirements", level=1)
    h1_hw.runs[0].font.color.rgb = RGBColor(31, 78, 120)
    doc.add_paragraph("Minimum hardware requirements for administrative, clinic, and staff usage.")

    min_bullets = [
        ("Type", "PC or Laptop (for administrative and clinic staff use)"),
        ("Processor", "Intel Core i3 or AMD Ryzen 3 (Dual-core 2.4GHz+)"),
        ("Memory (RAM)", "8GB DDR4 (Minimum) / 16GB (Recommended for dev)"),
        ("Storage", "256GB SSD (for fast system performance)"),
        ("Network", "Broadband Modem/Router with stable internet connection (Required for Gemini API)"),
        ("Peripherals", "High-resolution scanner/camera (for document uploads), standard 1080p monitor.")
    ]
    for label, val in min_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        r_lbl = bp.add_run(f"{label}: ")
        r_lbl.bold = True
        bp.add_run(val)

    doc.add_paragraph()

    # 4. SDLC Model: Agile Model
    h1_sdlc = doc.add_heading("4. SDLC Model: Agile Model", level=1)
    h1_sdlc.runs[0].font.color.rgb = RGBColor(31, 78, 120)
    doc.add_paragraph("The project follows the Agile SDLC model to allow for iterative development, continuous feedback, and flexibility in AI feature refinement.")

    doc.add_heading("Proposed Timetable (6 Months)", level=2).runs[0].font.color.rgb = RGBColor(46, 117, 182)
    agile_headers = ["Month", "Phase", "Key Activities"]
    agile_data = [
        ["Month 1", "Planning & Requirements", "Identifying system problems, objectives, and user needs."],
        ["Month 2", "System Design", "Architecture design, UI/UX prototyping, and AI logic planning."],
        ["Month 3", "Core Development", "Building Student Info Module and Database structure."],
        ["Month 4", "AI Integration", "AI implementation and ML tracking."],
        ["Month 5", "Testing & Validation", "System testing and bug fixing."],
        ["Month 6", "Deployment & Improvement", "Final deployment, staff training, and feature refinement."]
    ]
    create_styled_table(doc, agile_headers, agile_data, col_widths=[1.2, 2.3, 3.0])

    # 5. Features of the Proposed System
    h1_feat = doc.add_heading("5. Features of the Proposed System", level=1)
    h1_feat.runs[0].font.color.rgb = RGBColor(31, 78, 120)

    features = [
        ("Digital Enrollment Forms", "Streamlined student enrollment with document uploads."),
        ("Centralized Student Database", "Secure and organized storage for all student, academic, and administrative records."),
        ("AI-Powered Academic Monitoring", "Early Warning System (EWS) to detect declining performance trends based on grades."),
        ("Tuition Payment Risk Prediction", "Utilizes Machine Learning to classify payment risks and improve financial oversight."),
        ("Intelligent Report Generator", "Generates institutional and student reports in seconds using Gemini 2.0 Flash."),
        ("Responsive Student Portal", "Allows students and parents to monitor academic progress and records in real-time.")
    ]
    for title, desc in features:
        bp = doc.add_paragraph(style='List Bullet')
        r_t = bp.add_run(f"{title}: ")
        r_t.bold = True
        bp.add_run(desc)

    out_path = r"c:\Users\ender\Programming\Thesis_Project\Project_Documentation.docx"
    doc.save(out_path)
    print(f"Successfully updated {out_path}")

if __name__ == "__main__":
    generate_doc()
