import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F4E78") # Deep Blue
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)

    # Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F2F4F7" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            for p in row_cells[c_idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(50, 50, 50)

    # Set Column Widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph() # Spacing after table
    return table

def generate_doc():
    doc = Document()

    # Set Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Main Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_title = p_title.add_run("System Requirements & Tech Stack Documentation")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(31, 78, 120)

    # Subtitle
    p_sub = doc.add_paragraph()
    r_sub_lbl = p_sub.add_run("System: ")
    r_sub_lbl.bold = True
    r_sub_val = p_sub.add_run("Calvary Christian Academy (CCA) - Web-Based AI-Assisted Student Information and Academic Monitoring System")
    p_sub.runs[0].font.size = Pt(11)
    p_sub.runs[1].font.size = Pt(11)

    # Intro
    p_intro = doc.add_paragraph("This document outlines the updated technology stack, software prerequisites, and hardware requirements necessary to develop, deploy, and operate the system efficiently.")
    p_intro.runs[0].font.name = "Calibri"
    p_intro.runs[0].font.size = Pt(11)
    p_intro.paragraph_format.space_after = Pt(14)

    # Section 1
    h1 = doc.add_heading("1. Technology Stack", level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 120)
    
    p_stack = doc.add_paragraph("The application utilizes a modern decoupled architecture, combining a fast React-based frontend with a robust Python/FastAPI backend, integrated with Machine Learning and Generative AI capabilities.")
    p_stack.runs[0].font.name = "Calibri"
    p_stack.runs[0].font.size = Pt(11)

    # Frontend Table
    h2_fe = doc.add_heading("Frontend", level=2)
    h2_fe.runs[0].font.color.rgb = RGBColor(46, 117, 182)
    fe_headers = ["Component", "Technology", "Version / Detail"]
    fe_data = [
        ["Framework", "React", "^19.2.4"],
        ["Build Tool", "Vite", "^8.0.1"],
        ["Styling", "Tailwind CSS", "^4.2.2"],
        ["Language", "JavaScript (ES6+)", "With JSX"]
    ]
    create_styled_table(doc, fe_headers, fe_data, col_widths=[2.0, 2.0, 2.5])

    # Backend Table
    h2_be = doc.add_heading("Backend", level=2)
    h2_be.runs[0].font.color.rgb = RGBColor(46, 117, 182)
    be_headers = ["Component", "Technology", "Version / Detail"]
    be_data = [
        ["Framework", "FastAPI", "High-performance API framework"],
        ["ASGI Server", "Uvicorn", "[standard] for production-grade serving"],
        ["Database", "SQLite", "Development/Local deployment (.db files)"],
        ["ORM", "SQLAlchemy", "Relational database mapping"],
        ["Data Validation", "Pydantic", "Type checking and schema validation"],
        ["Authentication", "Passlib, python-jose", "Secure password hashing and JWT token management"]
    ]
    create_styled_table(doc, be_headers, be_data, col_widths=[2.0, 2.0, 2.5])

    # AI/ML Table
    h2_ai = doc.add_heading("AI & Machine Learning Integrations", level=2)
    h2_ai.runs[0].font.color.rgb = RGBColor(46, 117, 182)
    ai_headers = ["Component", "Technology", "Role"]
    ai_data = [
        ["Machine Learning", "Scikit-learn, NumPy", "RandomForest, GradientBoosting for predictive analytics (Early Warning Systems, Tuition Risk)."],
        ["Generative AI", "Google Generative AI", "Integration with Gemini (2.0 Flash) for automated narrative report generation."]
    ]
    create_styled_table(doc, ai_headers, ai_data, col_widths=[2.0, 2.0, 2.5])

    # Section 2
    h1_sw = doc.add_heading("2. Software Requirements", level=1)
    h1_sw.runs[0].font.color.rgb = RGBColor(31, 78, 120)
    p_sw = doc.add_paragraph("To run, develop, or deploy the application locally, the following software must be installed on the host machine:")
    p_sw.runs[0].font.name = "Calibri"

    sw_headers = ["Application", "Minimum Version", "Role"]
    sw_data = [
        ["Python", "3.10+", "Backend runtime environment. Required for FastAPI and ML models."],
        ["Node.js", "18.0+", "Frontend runtime environment. Required for Vite and npm packages."],
        ["npm", "9.0+", "Node package manager (comes with Node.js)."],
        ["Git", "Current Stable", "Version control system for repository management."],
        ["IDE / Code Editor", "VS Code / Cursor", "Recommended for development."],
        ["Web Browser", "Chrome, Edge, Safari", "Modern browser for testing and accessing the application."],
        ["API Testing", "Postman / Insomnia", "(Optional) For testing backend endpoints, though FastAPI provides built-in Swagger UI at /docs."]
    ]
    create_styled_table(doc, sw_headers, sw_data, col_widths=[1.8, 1.7, 3.0])

    # Section 3
    h1_hw = doc.add_heading("3. Hardware Requirements", level=1)
    h1_hw.runs[0].font.color.rgb = RGBColor(31, 78, 120)

    # Min Requirements
    h2_min = doc.add_heading("Minimum Requirements (For Staff/Administrative Use)", level=2)
    h2_min.runs[0].font.color.rgb = RGBColor(46, 117, 182)
    doc.add_paragraph("These are the minimum hardware specifications required for end-users (teachers, cashiers, registrars) accessing the web application.")
    
    min_bullets = [
        ("Processor", "Intel Core i3 or AMD Ryzen 3 (Dual-core 2.4GHz or higher)"),
        ("Memory (RAM)", "8GB DDR4"),
        ("Storage", "256GB Solid State Drive (SSD)"),
        ("Display", "Standard 1080p (1920x1080) monitor for optimal dashboard data visibility."),
        ("Peripherals", "Keyboard, Mouse, and an optional high-resolution scanner/camera (for physical document uploads).")
    ]
    for label, val in min_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        r_lbl = bp.add_run(f"{label}: ")
        r_lbl.bold = True
        r_val = bp.add_run(val)

    # Rec Requirements
    h2_rec = doc.add_heading("Recommended Requirements (For Local Development)", level=2)
    h2_rec.runs[0].font.color.rgb = RGBColor(46, 117, 182)
    doc.add_paragraph("Developers running both the frontend build process and backend machine learning models locally should meet these specifications.")

    rec_bullets = [
        ("Processor", "Intel Core i5 / AMD Ryzen 5 or equivalent (Quad-core or better)"),
        ("Memory (RAM)", "16GB (Running Vite, Uvicorn, VS Code, and ML models concurrently consumes significant memory)"),
        ("Storage", "512GB SSD")
    ]
    for label, val in rec_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        r_lbl = bp.add_run(f"{label}: ")
        r_lbl.bold = True
        r_val = bp.add_run(val)

    # Network
    h2_net = doc.add_heading("Network & Connectivity", level=2)
    h2_net.runs[0].font.color.rgb = RGBColor(46, 117, 182)
    bp_net = doc.add_paragraph(style='List Bullet')
    r_net_lbl = bp_net.add_run("Internet Connection: ")
    r_net_lbl.bold = True
    bp_net.add_run("A stable broadband connection is strictly required for both development and production environments. The system relies on external API calls to the ")
    r_gem = bp_net.add_run("Google Gemini API")
    r_gem.bold = True
    bp_net.add_run(" for report generation, which will fail if the system is offline.")

    # Save
    out_path = r"c:\Users\ender\Programming\Thesis_Project\SYSTEM_REQUIREMENTS.docx"
    doc.save(out_path)
    print(f"Successfully generated {out_path}")

if __name__ == "__main__":
    generate_doc()
