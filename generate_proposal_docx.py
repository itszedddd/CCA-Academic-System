import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    doc = docx.Document()

    # Title
    title = doc.add_heading('CCA Academic System — AI Feature Proposals', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Purpose
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Differentiate the Calvary Christian Academy (CCA) portal from generic school management systems by introducing unique, AI-driven features that leverage existing data.')

    doc.add_section()

    # Current AI Stack
    doc.add_heading('Current AI Stack', level=1)
    doc.add_paragraph('The CCA system already implements four AI-powered features:')

    # Current AI Stack Table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Feature'
    hdr_cells[2].text = 'Algorithm'
    hdr_cells[3].text = 'Data Source'
    hdr_cells[4].text = 'Purpose'

    data = [
        ['1', 'Grade Trend Analysis', 'Linear Regression', 'academic_records', 'Detects declining academic performance'],
        ['2', 'Resource Recommendations', 'Rule-based matching', 'academic_records', 'Suggests learning resources'],
        ['3', 'OCR Enrollment', 'Tesseract OCR', 'Uploaded images', 'Extracts text from forms'],
        ['4', 'Tuition Default Prediction', 'Payment ratio + multiplier', 'tuition_payments', 'Predicts risk of default']
    ]

    for item in data:
        row_cells = table.add_row().cells
        for i, text in enumerate(item):
            row_cells[i].text = text

    doc.add_paragraph()

    # Proposal 1
    doc.add_heading('Proposal 1: AI Attendance-Academic Correlation Engine (Recommended)', level=1)
    
    doc.add_heading('What It Does', level=2)
    doc.add_paragraph('Cross-references attendance patterns (absences, tardiness) with academic performance to generate early intervention alerts — before grades actually drop.')

    doc.add_heading('Why It\'s Unique', level=2)
    doc.add_paragraph('Most school management systems track attendance and grades as isolated modules. This feature connects them through a correlation algorithm, allowing the system to predict academic decline from behavioral signals.')

    doc.add_heading('Algorithm', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Pearson Correlation Coefficient between absence frequency and grade averages per student.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Sliding Window Analysis (2-week or 4-week) to detect recent behavioral shifts.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Threshold-based alerting: When correlation exceeds -0.6, generate an alert.')

    # Proposal 2
    doc.add_heading('Proposal 2: Smart Student Risk Profiling Dashboard', level=1)
    
    doc.add_heading('What It Does', level=2)
    doc.add_paragraph('Combines all existing AI signals (grade trend, tuition risk, attendance) into a unified composite risk score per student, visualized as a traffic-light dashboard.')

    doc.add_heading('Weighted Composite Score', level=2)
    p = doc.add_paragraph()
    p.add_run('composite_risk = (0.40 × academic_risk + 0.30 × financial_risk + 0.30 × attendance_risk)').italic = True

    # Risk Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sub-Score'
    hdr_cells[1].text = 'Source'
    hdr_cells[2].text = 'Normalization'

    risk_data = [
        ['academic_risk', 'Grade trend slope', 'max(0, min(1, slope / -10))'],
        ['financial_risk', 'Tuition risk_score', 'Already 0.0–0.95'],
        ['attendance_risk', 'Absence ratio', 'absences / total_days']
    ]

    for item in risk_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(item):
            row_cells[i].text = text

    # Proposal 3
    doc.add_heading('Proposal 3: AI Attendance Anomaly Detection', level=1)
    doc.add_heading('What It Does', level=2)
    doc.add_paragraph('Uses statistical anomaly detection to flag students whose attendance patterns have suddenly changed (behavioral shifts).')

    # Recommendation
    doc.add_heading('Recommendation', level=1)
    doc.add_paragraph('Go with Proposal 1 (Attendance-Academic Correlation) as the primary new feature. It is novel, algorithmically defensible, and requires zero schema changes. If time permits, layer Proposal 2 on top for a unified dashboard.')

    doc.save('CCA_AI_Feature_Proposals.docx')
    print("Document created successfully: CCA_AI_Feature_Proposals.docx")

if __name__ == "__main__":
    create_docx()
